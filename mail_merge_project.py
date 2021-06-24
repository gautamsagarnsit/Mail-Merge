#Use this command for python to exe:pyinstaller --onefile -w 'file_name.py'
import mammoth
import os
from bs4 import BeautifulSoup
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import smtplib,ssl
import docx2txt
import docx
import shutil
import win32com.client as win32
import sys




def login_user(email,password):
    sender_email = email
    password=password
    smtp_server = "smtp.gmail.com"
    port = 587  # For starttls 

    # Create a secure SSL context
    context = ssl.create_default_context()
    try:
        server = smtplib.SMTP(smtp_server,port)
        server.ehlo() # Can be omitted
        server.starttls(context=context) # Secure the connection
        server.ehlo() # Can be omitted
        server.login(sender_email, password)   
        msg="Email Credentials Validated"
        return msg,0,server,email 
    except Exception as e:
        # Print any error messages to stdout
        server.quit()
        msg="Login Unsuccessful: Check your account settings or Email/password"
        return msg,1


def send_merged_mail(Subject,source_name,docx_file_name,attachment_name,login_details):
    return_message=[]
    if len(login_details)!=4 or login_details[1]!=0:
        return_message.append("Login Failed")
        return return_message
    else:
        server=login_details[2]
    working_directory=os.getcwd()
    len_attachement=len(attachment_name) 

    return_message.append("No of attachments: "+str(len_attachement))  
    for i in range(len_attachement):
        return_message.append("Attachemnt "+str(i+1)+": "+str(attachment_name[i]))

    excel=win32.gencache.EnsureDispatch('Excel.Application')
    wb=excel.Workbooks.Open(source_name)
    ws=wb.Worksheets('Sheet1')

    def get_headers():
        i=1
        j=1
        placeholders=[]
        while ws.Cells(i,j).Value is not None:        
            while ws.Cells(i,j).Value is not None:
                if(i==1):
                    placeholders.append((ws.Cells(i,j).Value).upper())
                j+=1
            last_col=j-1
            j=1
            i+=1
        last_row=i-1
        return (last_row,last_col, placeholders)
    row,col,headers=get_headers()     

  
    doc=docx.Document(docx_file_name)
    src=docx_file_name
    dst=os.path.join(working_directory,"temp.docx") 
    shutil.copyfile(src, dst)
    temp=docx.Document(dst)
    merge_index=[]
    place_holders=[]
    para_index=[]
    Email_col=headers.index("Email".upper())+1
    i=0
    for para in temp.paragraphs:   
        for words in para.text.split():
            start=0
            end=0
            if "{{" and "}}" in words:
                iter=words.count("{")+words.count("}")
                iter=iter/4
                if(iter!=int(iter)):
                    print("Problem: Check the placement the Merge fields. Check if \"{{\" and \"}}\" are placed in right order")
                else:
                    iter=int(iter)
                for loop in range(iter):
                    start=words.find("{",end)
                    end=words.find("}",start)
                    place_holder=words[start+2:end]
                    
                    if(place_holder.upper() in headers):
                        place_holders.append(place_holder)
                        para_index.append(i)
                        index=headers.index(place_holder.upper())
                        merge_index.append(index+1)
                        return_message.append("\""+place_holder+"\""+": Merge field Added")
                    else:
                        print(place_holder,": Place holder not Exists")
                        return_message.append("\""+place_holder+"\""+": Merge field not found in source file")
        i+=1

    counter=0
    folder_name="Final Destination"
    destination_folder=os.path.join(working_directory,folder_name)
    while os.path.exists(destination_folder):
        counter+=1
        destination_folder=folder_name+"_"+str(counter)
    os.mkdir(destination_folder)
    folder_name="images"
    path_out=os.path.join(destination_folder,folder_name)

    return_message.append("Total Emails To Sent: "+str(row-1))
    return_message.append("Subject: "+Subject)

    for i in range(2,row+1):
        shutil.copyfile(src, dst)
        path_temp=os.path.abspath("temp.docx")
        temp=docx.Document(path_temp)
        for j in range(len(para_index)):
            temp.paragraphs[para_index[j]].text=temp.paragraphs[para_index[j]].text.replace("{{"+place_holders[j]+"}}",ws.Cells(i,merge_index[j]).Value)
            
        file_name="Schools template.docx"
        name="temp_"+str(i)+"_"+file_name    
        path=os.path.join(destination_folder,name)
        print(path)
        temp.save(path) 
        
        if  not os.path.exists(path_out):
                os.mkdir(path_out) 
                textresult=docx2txt.process(path,path_out)

        with open(path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value

        file_name ="docx_to_html.html"
        file=open(os.path.join(destination_folder,file_name),"w")
        file.write(html)
        file.close()

            # instance of MIMEMultipart
        msg = MIMEMultipart()
        sender_email=login_details[3]
        # storing the senders email address  
        msg['From'] = sender_email
        
        # storing the receivers email address 
        msg['To'] = ws.Cells(i,Email_col).Value
        
        # storing the subject 
        msg['Subject'] = Subject

        soup=BeautifulSoup()
        

        file=open(os.path.join(destination_folder,file_name),'r')
        soup=BeautifulSoup(file,'html.parser')
        result=soup.find_all('img')
        image_files=os.listdir(path_out)

        img_index=0
        for img in result:
            img_name=image_files[img_index]
            img_index+=1
            path=os.path.join(path_out,img_name)
            img['src']="cid:image"+str(img_index)
            fp = open(path, 'rb')
            msgImage = MIMEImage(fp.read())
            # Do not forget close the file object after using it.
            fp.close()
            msgImage.add_header('Content-ID', '<image'+str(img_index)+ '>')
            # Attach the MIMEImage object to the email body.
            msg.attach(msgImage)
        msgtext=MIMEText(str(soup), 'html')
        msg.attach(msgtext)
        file.close()

        for attach_index in range(len(attachment_name)):  
            # open the file to be sent 
            attachment_file = attachment_name[attach_index]
            print(attachment_file)
            attachment = open(os.path.normpath(attachment_file)[1:-1], "rb")

            # instance of MIMEBase and named as p
            p=(MIMEBase('application', 'octet-stream'))

            # To change the payload into encoded form
            p.set_payload((attachment).read())

            # encode into base64
            encoders.encode_base64(p)

            p.add_header('Content-Disposition', "attachment; filename= %s" % os.path.basename(attachment_file)[:-1])

            msg.attach(p)        

        # Converts the Multipart msg into a string
        text = msg.as_string()

        if ws.Cells(i,Email_col).Value==None:
            print("No Email Address Found")
        else:
            receiver_email=ws.Cells(i,Email_col).Value
            
            try:
                server.sendmail(sender_email, receiver_email, text) 
                print("Email Sent "+str(i))
                return_message.append("Email Sent to "+receiver_email)  
            except:
                print("Email Not Sent to "+receiver_email)
                return_message.append("Email Not Sent to "+receiver_email)
                

    os.remove(dst)
    shutil.rmtree(destination_folder,ignore_errors=True)
    server.quit()   
    wb.Close(False)
    return_message.append("Mail Merge Completed!")
    return return_message
